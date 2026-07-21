# -*- coding: utf-8 -*-
"""Model Validasyonu - Archer field-config Word teknik build notu."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x54,0x96); GREY=RGBColor(0x59,0x59,0x59)
doc=Document()
nm=doc.styles["Normal"]; nm.font.name="Calibri"; nm.font.size=Pt(10.5)
nm.paragraph_format.space_after=Pt(6); nm.paragraph_format.line_spacing=1.12
for lvl,sz,col in [("Heading 1",16,NAVY),("Heading 2",13,BLUE),("Heading 3",11.5,BLUE)]:
    s=doc.styles[lvl]; s.font.name="Calibri"; s.font.size=Pt(sz); s.font.color.rgb=col; s.font.bold=True
def shade(cell,hx):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:color"),"auto"); sh.set(qn("w:fill"),hx); tcPr.append(sh)
def setc(cell,t,bold=False,color=None,size=9.5,fill=None,mono=False):
    cell.text=""; p=cell.paragraphs[0]; run=p.add_run(str(t))
    run.font.size=Pt(size); run.font.bold=bold; run.font.name="Consolas" if mono else "Calibri"
    if color: run.font.color.rgb=color
    if fill: shade(cell,fill)
def tbl(headers,rows,widths=None,font=9,mono_cols=()):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers): setc(t.rows[0].cells[j],h,bold=True,color=RGBColor(0xFF,0xFF,0xFF),size=font,fill="1F3864")
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,val in enumerate(row): setc(cells[j],val,size=font,fill=("F2F2F2" if i%2 else None),mono=(j in mono_cols))
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Cm(w)
    return t
def para(t,italic=False,size=10.5,color=None):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=italic; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    return p
def bullet(t,lead=None):
    p=doc.add_paragraph(style="List Bullet")
    if lead: r=p.add_run(lead); r.bold=True; p.add_run(" – "+t)
    else: p.add_run(t)
    return p

tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("MODEL VALİDASYONU"); r.font.size=Pt(26); r.font.bold=True; r.font.color.rgb=NAVY
sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp.add_run("Archer Field-Config – Teknik Kurulum Notu"); r.font.size=Pt(15); r.font.bold=True; r.font.color.rgb=BLUE
for line in ["Hedef birim: Yazılım Geliştirme Daire Başkanlığı",
             "Ek: Model_Validasyonu_Archer_Field_Config.xlsx (13 sekme, 60 alan)","Taslak v0.1"]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=p.add_run(line); rr.font.size=Pt(11); rr.font.color.rgb=GREY
doc.add_paragraph()
para("Bu not, Model Validasyonu uygulamasının Archer üzerinde kurulumunu anlatır; alan-seviyesi tam konfigürasyon ekteki "
     "Excel’dedir. Metodolojik omurga: SR 11-7 üç bileşen (kavramsal sağlamlık · süreç doğrulama · çıktı analizi), "
     "PRA SS1/23 Principle 4 (4.1–4.5) ve ECB TRIM validasyon test kataloğu; Archer MRM App-Pack 'Model Validation' "
     "yapısıyla hizalıdır.", italic=True, size=9.5, color=GREY)
doc.add_page_break()

doc.add_heading("1. Uygulama Künyesi", level=1)
tbl(["Özellik","Değer"],[
 ["Uygulama adı","Model Validasyonu (Model Validation)"],
 ["Workspace","Model Riski Yönetimi (MRM)"],
 ["Anahtar alan","Validasyon Başlığı (Text). Tekil kimlik: Validasyon ID (Tracking ID)."],
 ["Ana ilişki","İlgili Model → Model Envanteri"],
 ["İş akışı","Planlandı → Devam → Bulgu Gözden Geçirme → Onayda → Tamamlandı"],
 ["Alan sayısı","60 (Kimlik/Bağımsızlık, Kavramsal Sağlamlık, Süreç Doğrulama, Çıktı Analizi, Sonuç, Onay, Takip, Sistem)"],
],widths=[4.5,13.5],font=9.5)

doc.add_heading("2. Temel Tasarım İlkesi: ÖNERİ ≠ KARAR", level=1)
para("SS1/23 4.1b ve 2.13 uyarınca bağımsız validasyon fonksiyonu bir ÖNERİ üretir; ONAY yetkisi ayrı bir onay makamındadır "
     "(Model Risk Komitesi / üst yönetim). Bu nedenle uygulamada:")
bullet("[Validasyon Sonucu (ÖNERİ)] ve [Onay Kararı] AYRI iki alandır.", lead="İki ayrı alan")
bullet("Öneri alanları yalnız 2. hat bağımsız validasyona; onay alanları yalnız Model Risk Komitesine yazılabilir (FLS).", lead="İki ayrı rol")
bullet("Advanced Workflow’da öneri (N4) ile onay (N5) ayrı node’lardır.", lead="İki ayrı adım")

doc.add_heading("3. Alan Grupları (Özet)", level=1)
for g,d in [
 ("Kimlik (1-9)","Validasyon ID/başlık, ilgili model, getirilen tier, validasyon türü, gerekçe/tetik, kapsam, planlı/fiili tarihler"),
 ("Bağımsızlık (10-13)","Validatör, validasyon birimi, bağımsızlık teyidi (ECB 3 opsiyon + dış uzman), dış uzman kuruluş"),
 ("Kavramsal Sağlamlık (14-21)","Teori/metodoloji, amaca uygunluk, varsayım, duyarlılık, veri kalitesi/temsil, geliştirme kanıtı, nitel değerlendirme, sonuç"),
 ("Süreç Doğrulama (22-27)","Girdi/hesaplama-kod QA/entegrasyon/çıktı doğrulama, IT birebir üretim teyidi, sonuç"),
 ("Çıktı Analizi & Testler (28-37)","Backtesting, ayrım gücü, temsil, override, stabilite, benchmarking, paralel çıktı, eşik ihlali, metrik linki, ilk-only görevler"),
 ("Sonuç ve Bulgular (38-46)","Etkin sorgulama, sınırlamalar, üretilen bulgular, ÖNERİ, rating, koşullar, rapor/tarih"),
 ("Onay (47-52)","Onaya sunum, onay makamı, ONAY KARARI, tarih, koşullar, üst yönetime raporlama"),
 ("Takip (53-55)","Sonraki revalidasyon tarihi (hesaplanan), açık aksiyonlar, validasyon durumu"),
 ("Sistem (56-60)","Kayıt izinleri, tarihler, History Log"),
]:
    bullet(d, lead=g)

doc.add_heading("4. Validasyon Test Kataloğu", level=1)
para("Çıktı analizi bölümü, ECB TRIM ve SR 11-7 test setine göre kurulmuştur; testlerin sıklığı (ilk / yıllık / periyodik / "
     "ilk-only) ve alan eşlemesi Excel “5. Validasyon Test Kataloğu” sekmesindedir. Öne çıkanlar:")
bullet("Her validasyon: kavramsal sağlamlık + süreç doğrulama + nitel analizler + eşik değerlendirmesi.")
bullet("En az yıllık: backtesting, ayrım gücü (AUC/Gini/KS), temsil, override, stabilite analizleri (tier’e göre yoğunluk).")
bullet("Periyodik: benchmarking / kıyaslama.")
bullet("Yalnız ilk validasyon + materyal değişiklik: model geliştirme replikasyonu, kod QA, IT birebir üretim teyidi.")

doc.add_heading("5. Kritik Mekanizmalar", level=1)
doc.add_heading("5.1 Hesaplanan alanlar", level=2)
para("Sonraki revalidasyon tarihi, rapor tarihine modelin validasyon frekansının (ay) eklenmesiyle hesaplanır:")
p=doc.add_paragraph(); r=p.add_run('IF(ISEMPTY([Rapor Tarihi]),"",DATEADD([Rapor Tarihi],"m",SELECTEDVALUENUMBER(MOSTRECENTVALUE(REF([İlgili Model],[Validasyon Frekansı])))))')
r.font.name="Consolas"; r.font.size=Pt(8.5)
bullet("Model Tier getirilen alanı: MOSTRECENTVALUE(REF([İlgili Model],[Tier])) — kapsam/yoğunluğu yönlendirir.")
bullet("Süre (gün) = DATEDIF(Fiili Başlangıç, Fiili Bitiş); Bulgu Sayısı ilişkili bulgu adedi.")

doc.add_heading("5.2 DDE (koşullu davranış)", level=2)
bullet("Validasyon Türü = İlk → replikasyon/kod QA/IT uygunluk alanları görünür ve zorunlu.")
bullet("Model Tier = Tier 1 → backtesting/ayrım gücü/derin kapsam alanları zorunlu.")
bullet("Bağımsızlık = Dış Uzman → Dış Uzman Kuruluş (Tedarikçi) zorunlu.")
bullet("Eşik İhlali = İhlal → bulgu oluştur uyarısı + Üretilen Bulgular öne çıkar.")
bullet("Sonuç = Koşullu Onay Önerilir → Önerilen Koşullar zorunlu; Onay = Koşullu → Onay Koşulları zorunlu.")

doc.add_heading("5.3 Advanced Workflow", level=2)
para("Planlandı → Devam Ediyor → Bulgu Gözden Geçirme → (öneri damgalanır) → Onayda → Onayla/Koşullu/Reddet → Tamamlandı. "
     "Onay node’unda Model Envanteri otomatik güncellenir (durum + son validasyon tarihi + sonraki revalidasyon tarihi).")

doc.add_heading("5.4 Bağımsızlık (FLS)", level=2)
para("1. hat (geliştirici/sahip) validasyon kaydının Etkin Sorgulama / Sonuç / ÖNERİ bölümlerini göremez/değiştiremez; bu "
     "bölümler yalnız bağımsız validasyona açıktır. Onay alanları yalnız Model Risk Komitesine yazılabilir. Detay: Excel "
     "“9. Erişim & FLS”. Bu düzen SR 11-7 / SS1/23 / ECB bağımsızlık ilkesinin sistemsel karşılığıdır.")

doc.add_heading("6. Mevcut Uygulamalarla Bağlantı", level=1)
tbl(["Validasyon alanı","Bağlanan mevcut uygulama"],[
 ["İlgili Model","Model Envanteri (ters: Validasyon Kayıtları)"],
 ["Üretilen Bulgular","Bulgular (=Issues Management)"],
 ["Açık Aksiyonlar / Remediation","Aksiyon Planları"],
 ["İlgili İzleme Metrikleri","Metrik Sonuçları"],
 ["Dış Uzman Kuruluş","Tedarikçi (BDDK 'uzman kuruluş')"],
 ["Validasyon Birimi","Organizasyon / İş Hiyerarşisi (2. hat)"],
],widths=[7.0,11.0],font=9.5)

doc.add_heading("7. Doğrulanacak Noktalar", level=1)
bullet("DATEADD ay-jetonu ('m') ve referanslı Values List üzerinde SELECTEDVALUENUMBER kullanımı canlı Archer örneğinde teyit edilmeli (alternatif: frekansı sayısal alan olarak REF ile bu kayda çekip DATEADD uygulamak).")
bullet("Bulgu sayısı için REF/COUNT davranışı canlı örnekte doğrulanmalı; gerekirse rollup alanı kullanılmalı.")
bullet("App-Pack lisans durumu → alan adları App-Pack 'Model Validation' veri sözlüğüyle hizalanmalı.")
bullet("Tier→validasyon kapsam/frekans eşiği ve test eşikleri (threshold) banka MRM politikasıyla kalibre edilmeli.")
bullet("BDDK 'uzman kuruluş' ile dış validasyon senaryosunda sözleşme/gizlilik gereklilikleri hukukla teyit edilmeli.")

para("")
para("Not: Bu, regülasyon ve Archer platform araştırmasına dayalı eksiksiz bir başlangıç konfigürasyonudur; banka Archer "
     "sürümü, lisanslı modüller ve mevcut uygulama şemalarıyla birlikte nihai hale getirilecektir.", italic=True, size=9.5, color=GREY)

out="/home/user/rrr/mrm-archer/Model_Validasyonu_Archer_Field_Config_Not.docx"
doc.save(out); print("KAYDEDILDI:",out)
