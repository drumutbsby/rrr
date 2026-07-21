# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x54,0x96); GREY=RGBColor(0x59,0x59,0x59)
doc=Document()
nm=doc.styles["Normal"]; nm.font.name="Calibri"; nm.font.size=Pt(10.5); nm.paragraph_format.space_after=Pt(6); nm.paragraph_format.line_spacing=1.12
for lvl,sz,col in [("Heading 1",16,NAVY),("Heading 2",13,BLUE)]:
    s=doc.styles[lvl]; s.font.name="Calibri"; s.font.size=Pt(sz); s.font.color.rgb=col; s.font.bold=True
def shade(cell,hx):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:color"),"auto"); sh.set(qn("w:fill"),hx); tcPr.append(sh)
def setc(cell,t,bold=False,color=None,size=9.5,fill=None,mono=False):
    cell.text=""; p=cell.paragraphs[0]; run=p.add_run(str(t)); run.font.size=Pt(size); run.font.bold=bold; run.font.name="Consolas" if mono else "Calibri"
    if color: run.font.color.rgb=color
    if fill: shade(cell,fill)
def tbl(headers,rows,widths=None,font=9,mono_cols=()):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers): setc(t.rows[0].cells[j],h,bold=True,color=RGBColor(0xFF,0xFF,0xFF),size=font,fill="1F3864")
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,val in enumerate(row): setc(cells[j],val,size=font,fill=("F2F2F2" if i%2 else None),mono=(j in mono_cols))
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Cm(w)
    return t
def para(t,italic=False,size=10.5,color=None):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=italic; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    return p
def bullet(t,lead=None):
    p=doc.add_paragraph(style="List Bullet")
    if lead: r=p.add_run(lead); r.bold=True; p.add_run(" – "+t)
    else: p.add_run(t)
    return p
tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("MODEL KRİTİKLİK DEĞERLENDİRMESİ"); r.font.size=Pt(23); r.font.bold=True; r.font.color.rgb=NAVY
sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp.add_run("Archer Questionnaire Field-Config – Teknik Kurulum Notu"); r.font.size=Pt(14); r.font.bold=True; r.font.color.rgb=BLUE
for line in ["Hedef birim: Yazılım Geliştirme Daire Başkanlığı","Ek: Model_Kritiklik_Degerlendirmesi_Archer_Field_Config.xlsx (12 sekme, 35 soru/alan)","Taslak v0.1"]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=p.add_run(line); rr.font.size=Pt(11); rr.font.color.rgb=GREY
doc.add_paragraph()
para("Bu, bir Archer QUESTIONNAIRE (anket) uygulamasıdır; standart uygulamadan farklı olarak alanları 'soru'dur ve "
     "WEIGHTEDSCORE fonksiyonu yalnız burada çalışır. Materyalite × Karmaşıklık iki boyutundan İçsel Risk skoru ve Tier "
     "üretir; sonuç, Model Envanteri'ne REF ile geri yazılır. Kaynak: PRA SS1/23 §1.3, OSFI E-23 risk rating.", italic=True, size=9.5, color=GREY)
doc.add_page_break()
doc.add_heading("1. Skorlama Modeli", level=1)
para("İki boyut, her blokta ağırlıklar %100'e normalize; her cevap 1-5 sayısal değer taşır. WEIGHTEDSCORE = cevap değeri × soru ağırlığı.")
tbl(["Boyut","Skor","Formül (özet)"],[
 ["Materyalite (1-5)","M1..M5 ağırlıklı","Σ WEIGHTEDSCORE(M1..M5), ağırlıklar Σ=1"],
 ["Karmaşıklık (1-5)","K1..K7 ağırlıklı (AI faktörleri dahil)","Σ WEIGHTEDSCORE(K1..K7), ağırlıklar Σ=1"],
 ["İçsel Risk (1-25)","Materyalite × Karmaşıklık","ROUND(Mat×Karm,2)"],
 ["Tier","İçsel Risk eşiği + sert kurallar","≥15→T1, ≥8→T2, ≥3→T3, <3→İhmal"],
],widths=[5.0,6.5,6.5],font=9.5)
doc.add_heading("2. Sert Kurallar (tier eskalasyonu)", level=1)
bullet("Düzenleyici/sermaye kullanımı (IRB, İSEDES, TFRS9, piyasa/AMA) = Evet → minimum Tier 2; 'İhmal Edilebilir' SEÇİLEMEZ (DDE ile listeden filtrelenir).")
bullet("Tam otonom karar (O1=3) + yüksek müşteri etkisi → minimum Tier 2.")
bullet("Çok sayıda downstream bağımlılık (B1=3) → önerilen tier bir kademe yükselir.")
bullet("GenAI/Temel Model = Evet → minimum Tier 2 + AI/ML ek değerlendirmesi zorunlu.")
doc.add_heading("3. Questionnaire Mekaniği & Geri Yazım", level=1)
bullet("Her skorlu soru bir Values List; SORU AĞIRLIĞI (weighting) ve CEVAP başına SAYISAL DEĞER taşır.")
bullet("AI/ML modelinde K5 (açıklanabilirlik), K6 (yanlılık), O1 (otonomi) DDE ile zorunlu olur.")
bullet("2. hat bağımsız validasyon Nihai Tier'i teyit eder / override edebilir (gerekçe zorunlu) — SS1/23 1.3e.")
bullet("Model Envanteri, bu anketten Materyalite/Karmaşıklık/İçsel Risk/Tier alanlarını MOSTRECENTVALUE(REF(...)) ile çeker.")
doc.add_heading("4. Doğrulanacak Noktalar", level=1)
bullet("Her M/K sorusu için 1-5 ölçeğine banka risk iştahına göre sayısal 'anchor' cümleleri yazılmalı (ör. M1 tutar bantları).")
bullet("Blok ağırlıkları ve tier eşikleri (15/8/3) kalibre edilmeli; sert kurallar hukuk/risk ile teyit edilmeli.")
bullet("WEIGHTEDSCORE questionnaire-only kısıtı ve referanslı VL üzerinde SELECTEDVALUENUMBER kullanımı canlı örnekte doğrulanmalı.")
para("")
para("Not: Regülasyon ve Archer platform araştırmasına dayalı başlangıç konfigürasyonudur; banka risk iştahı ve Archer sürümüyle nihai hale getirilecektir.", italic=True, size=9.5, color=GREY)
out="/home/user/rrr/mrm-archer/Model_Kritiklik_Degerlendirmesi_Archer_Field_Config_Not.docx"
doc.save(out); print("KAYDEDILDI:",out)
