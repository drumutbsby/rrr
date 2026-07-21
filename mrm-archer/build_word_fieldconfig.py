# -*- coding: utf-8 -*-
"""Model Envanteri - Archer field-config Word teknik build notu."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x54,0x96); GREY=RGBColor(0x59,0x59,0x59)
doc=Document()
n=doc.styles["Normal"]; n.font.name="Calibri"; n.font.size=Pt(10.5)
n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.12
for lvl,sz,col in [("Heading 1",16,NAVY),("Heading 2",13,BLUE),("Heading 3",11.5,BLUE)]:
    s=doc.styles[lvl]; s.font.name="Calibri"; s.font.size=Pt(sz); s.font.color.rgb=col; s.font.bold=True

def shade(cell,hx):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:color"),"auto"); sh.set(qn("w:fill"),hx); tcPr.append(sh)
def setc(cell,t,bold=False,color=None,size=9.5,fill=None,mono=False):
    cell.text=""; p=cell.paragraphs[0]; run=p.add_run(str(t))
    run.font.size=Pt(size); run.font.bold=bold; run.font.name="Consolas" if mono else "Calibri"
    if color: run.font.color.rgb=color
    if fill: shade(cell,fill)
def tbl(headers,rows,widths=None,font=9,mono_cols=()):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers): setc(t.rows[0].cells[j],h,bold=True,color=RGBColor(0xFF,0xFF,0xFF),size=font,fill="1F3864")
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,val in enumerate(row): setc(cells[j],val,size=font,fill=("F2F2F2" if i%2 else None),mono=(j in mono_cols))
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Cm(w)
    return t
def para(t,italic=False,size=10.5,color=None):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=italic; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    return p
def bullet(t,lead=None):
    p=doc.add_paragraph(style="List Bullet")
    if lead: r=p.add_run(lead); r.bold=True; p.add_run(" – "+t)
    else: p.add_run(t)
    return p

# Kapak
tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("MODEL ENVANTERİ"); r.font.size=Pt(26); r.font.bold=True; r.font.color.rgb=NAVY
sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp.add_run("Archer Field-Config – Teknik Kurulum Notu"); r.font.size=Pt(15); r.font.bold=True; r.font.color.rgb=BLUE
for line in ["Hedef birim: Yazılım Geliştirme Daire Başkanlığı",
             "Ek: Model_Envanteri_Archer_Field_Config.xlsx (12 sekme, 79 alan)",
             "Taslak v0.1"]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=p.add_run(line); rr.font.size=Pt(11); rr.font.color.rgb=GREY
doc.add_paragraph()
para("Bu not, Model Envanteri uygulamasının Archer üzerinde nasıl kurulacağını anlatır; alan-seviyesi tam konfigürasyon "
     "(her alanın Archer tipi, zorunluluğu, seçim listesi, cross-reference hedefi, hesaplama formülü, DDE davranışı), ekteki "
     "Excel çalışma kitabındadır. Tasarım, OSFI E-23 asgari envanter şeması + PRA SS1/23 §1.2c + SR 11-7 + Archer MRM "
     "App-Pack yapısı çapraz doğrulamasına dayanır.", italic=True, size=9.5, color=GREY)
doc.add_page_break()

doc.add_heading("1. Uygulama Künyesi", level=1)
tbl(["Özellik","Değer"],[
 ["Uygulama adı","Model Envanteri (Model Inventory)"],
 ["Workspace","Model Riski Yönetimi (MRM) – yeni"],
 ["Uygulama tipi","Standart (leveled değil); kayıt-seviyesi güvenlik açık"],
 ["Anahtar alan","Model Adı (Text, unique). Tekil sistem kimliği: Model ID (Tracking ID)."],
 ["İş akışı","Advanced Workflow (yaşam döngüsü + validasyon/onay + sertifikasyon)"],
 ["Alan sayısı","79 (12 mantıksal grup, 6 kullanıcı sekmesi + Sistem)"],
],widths=[4.5,13.5],font=9.5)

doc.add_heading("2. Tasarım Kararları", level=1)
doc.add_heading("2.1 App-Pack mı, özel kurulum mu?", level=2)
para("Archer, hazır bir “Model Risk Management App-Pack” sunar (Platform 6.10+; Business Unit ve Issues Management "
     "kullanım senaryolarını gerektirir). Bunun Model Inventory uygulaması General Information alanları — Model Name, "
     "Model Version, Model Type, Model Category, Priority, Model Implementation, Purpose — ve Business Unit Owner / Model "
     "Risk Program Manager rolleri, bu tasarımla birebir örtüşür.")
bullet("Banka App-Pack’i lisanslıyorsa: yeni Model Envanteri/Validasyon/Değişiklik/Sertifikasyon yapısını App-Pack tabanına hizalayın; Business Unit ve Issues Management cross-reference’larını yeniden kullanın; paralel yapı kurmayın.", lead="Senaryo A")
bullet("Lisans yoksa: aynı yapıyı core Archer alan tipleriyle kurun (bu Excel tam bunu tarifler). Her iki senaryoya da uygundur.", lead="Senaryo B")

doc.add_heading("2.2 Mevcut uygulamaların yeniden kullanımı", level=2)
tbl(["Model Envanteri alanı","Mevcut Archer uygulaması","Not"],[
 ["Sahip İş Birimi","Organizasyon / İş Hiyerarşisi","= Archer 'Business Unit'"],
 ["Açık Bulgular / Aksiyon Planları","Bulgular + Aksiyon Planları","= 'Issues Management'; yeni bulgu uygulaması AÇILMAZ"],
 ["Barındıran Sistem / Donanım","Uygulamalar / Donanımlar / IT Hizmetleri","BT altyapısı bağı"],
 ["Girdi Veri Varlıkları","Bilgi Varlıkları Envanteri","Veri/lineage"],
 ["Telafi Edici Kontroller","Kontroller","Model riskini azaltan kontroller"],
 ["İzleme","Metrikler / Metrik Sonuçları","MRM metrik seti; eşik aşımı → Bulgu"],
 ["İlgili Mevzuat / Politika","Yasal Mevzuatlar / Politikalar","BDDK eşleme ve MRM politikası"],
],widths=[5.5,6.5,6.0],font=9)

doc.add_heading("3. Alan Grupları (Özet)", level=1)
para("Tam alan listesi Excel “1. Field Config” sekmesindedir. 79 alan, kaynağıyla birlikte şu gruplara ayrılır:")
for g,d in [
 ("Kimlik (1-4)","Model ID (Tracking ID), Model Adı (key/unique), tanım, versiyon"),
 ("Sınıflandırma (5-11)","Tip, kategori (çoklu), öncelik, AI/ML·GenAI·Dinamik·Deterministik bayrakları"),
 ("Amaç ve Kullanım (12-17)","Amaç, ürün/portföy, onaylı vs fiili kullanım, kısıtlar, operasyonel sınırlar"),
 ("Sahiplik ve Roller (18-24)","Sahip birim, owner, developer, user, validator, approver, MRM programı"),
 ("Köken ve Barındırma (25-29)","Köken, tedarikçi, grup validasyonu, barındıran sistem/donanım"),
 ("Kritiklik ve Risk (30-36)","Kritiklik anketi linki, materyalite/karmaşıklık/içsel risk skorları, tier, risk derecesi, artık risk"),
 ("Metodoloji ve Veri (37-46)","Metodoloji, hedef değişken, algoritma, veri kaynakları/varlıkları, çıktı, upstream/downstream, dokümantasyon"),
 ("Varsayım-Sınır-Bulgu (47-56)","Varsayım, sınırlama, yakalanamayan risk, telafi kontrolleri, sağlık göstergesi, açık bulgular, mevzuat/politika"),
 ("Yaşam Döngüsü ve Takvim (57-71)","Durum, tarihler, frekans, sonraki validasyon (hesaplanan), izleme/sertifikasyon durumu, ilişkili kayıtlar"),
 ("Emeklilik (72-74)","Gerekçe, tarih, saklama süresi"),
 ("Sistem (75-79)","Kayıt izinleri, tarihler, History Log"),
]:
    bullet(d, lead=g)

doc.add_heading("4. Kritik Mekanizmalar", level=1)
doc.add_heading("4.1 Hesaplanan alanlar (Archer formülü)", level=2)
para("Sonraki validasyon tarihi, frekansın (ay olarak sayısal eşlenmiş) son validasyon tarihine eklenmesiyle hesaplanır:")
p=doc.add_paragraph(); r=p.add_run('IF(ISEMPTY([Son Validasyon Tarihi]),"",DATEADD([Son Validasyon Tarihi],"m",SELECTEDVALUENUMBER([Validasyon Frekansı])))')
r.font.name="Consolas"; r.font.size=Pt(9)
bullet("Gecikme bayrağı: sonraki tarih geçmişse 'Gecikmiş', 30 gün içindeyse 'Yaklaşan', aksi 'Zamanında'.")
bullet("Materyalite/karmaşıklık skorları: Kritiklik questionnaire'inden WEIGHTEDSCORE ile hesaplanıp REF/MOSTRECENTVALUE ile envantere çekilir (WEIGHTEDSCORE yalnız questionnaire'de çalışır).")
bullet("İçsel risk skoru = materyalite × karmaşıklık; tier bu skordan politika eşikleriyle önerilir, validasyonda teyit edilir.")

doc.add_heading("4.2 Data-Driven Events (koşullu davranış)", level=2)
bullet("AI/ML bayrağı = Evet → AI alanları ve AI/ML Ek anketi görünür/zorunlu.")
bullet("Deterministik/EUC = Evet → validasyon sekmeleri gizlenir, 'kontrol-yalnızca' yol açılır (SS1/23 1.1c).")
bullet("Köken = Satıcı → Tedarikçi zorunlu; Köken = Grup-Ana → grup validasyon koşulları zorunlu.")
bullet("Tier ∈ {1,2} → OSFI 'non-negligible' ek alanları (versiyon, operasyonel sınır, validatör, veri kaynağı) zorunlu.")
bullet("Durum = Koşullu Onaylı → PMA/İstisna zorunlu; Durum = Emekli → emeklilik gerekçesi/tarihi zorunlu.")

doc.add_heading("4.3 Advanced Workflow (yaşam döngüsü)", level=2)
para("Taslak → Geliştirme → Bağımsız Validasyonda → (Evaluate) → Onayda → Onaylı/Prod (veya Koşullu Onaylı) → İzlemede "
     "→ (Revalidasyon tetiği) → … → Emekli. User Action node’larının geçişleri Approve/Reject/Return butonları olur; "
     "Update Content node’ları Model Durumu’nu damgalar. Onay kararı (Model Risk Komitesi) validasyon önerisinden ayrıdır.")

doc.add_heading("4.4 Bağımsızlık (alan-seviyesi güvenlik)", level=2)
para("Bağımsız validasyon fonksiyonu, geliştirme/sahiplik alanlarını salt-okur görür; skor/tier teyidi ve validasyon "
     "kayıtlarında yazma yetkisine sahiptir. Bu FLS düzeni, SR 11-7 / SS1/23 / ECB bağımsızlık ilkesinin sistemsel "
     "karşılığıdır (Excel “8. Erişim & FLS”).")

doc.add_heading("5. Doğrulanacak Noktalar", level=1)
bullet("DATEADD ay-kaydırma jetonu ('m' / 'month') canlı Archer örneğinde veya fonksiyon referansında birebir teyit edilmeli.")
bullet("Numeric alan min/max/ondalık özellik etiketleri canlı örnekte doğrulanmalı.")
bullet("App-Pack lisans durumu netleşince, alan adları App-Pack veri sözlüğüyle birebir hizalanmalı.")
bullet("Tier→frekans eşiği (P1/P2/P3) ve validasyon frekansları banka MRM politikası/risk iştahıyla kalibre edilmeli.")

para("")
para("Not: Bu, regülasyon ve Archer platform araştırmasına dayalı eksiksiz bir başlangıç konfigürasyonudur; banka Archer "
     "sürümü, lisanslı modüller ve mevcut uygulama şemalarıyla birlikte nihai hale getirilecektir.", italic=True, size=9.5, color=GREY)

out="/home/user/rrr/mrm-archer/Model_Envanteri_Archer_Field_Config_Not.docx"
doc.save(out); print("KAYDEDILDI:",out)
