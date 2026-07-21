# -*- coding: utf-8 -*-
"""MRM - Archer Entegrasyonu: Word is analizi / cozum tasarimi dokumani."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x54, 0x96)
GREY = RGBColor(0x59, 0x59, 0x59)

doc = Document()

# ---- Base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for lvl, sz, col in [("Heading 1", 16, NAVY), ("Heading 2", 13, BLUE), ("Heading 3", 11.5, BLUE)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)

def set_cell(cell, text, bold=False, color=None, size=9.5, align="left", fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    run = p.add_run(str(text))
    run.font.size = Pt(size); run.font.bold = bold; run.font.name = "Calibri"
    if color: run.font.color.rgb = color
    if fill: shade(cell, fill)

def add_table(headers, rows, widths=None, header_fill="1F3864", font=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        set_cell(t.rows[0].cells[j], h, bold=True, color=RGBColor(0xFF,0xFF,0xFF), size=font, fill=header_fill)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            set_cell(cells[j], val, size=font, fill=("F2F2F2" if i%2 else None))
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Cm(w)
    return t

def bullet(text, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet" if level==0 else "List Bullet 2")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True; r.font.size = Pt(10.5)
        p.add_run(" – " + text).font.size = Pt(10.5)
    else:
        p.add_run(text).font.size = Pt(10.5)
    return p

def para(text, italic=False, size=10.5, color=None, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    return p

# =========================================================
# KAPAK
# =========================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("MODEL RİSKİ YÖNETİMİ (MRM)")
r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = NAVY
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Archer Entegrasyonu – İş Analizi ve Çözüm Tasarımı Dokümanı")
r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = BLUE
doc.add_paragraph()
for line, sz, col, bold in [
    ("Hedef birim: Yazılım Geliştirme Daire Başkanlığı", 12, GREY, True),
    ("Doküman türü: İş analizi / gereklilik ve çözüm tasarımı taslağı", 11, GREY, False),
    ("Statü: Taslak v0.1 — regülasyon araştırmasına dayalı, gözden geçirme bekliyor", 11, GREY, False),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run(line); rr.font.size = Pt(sz); rr.font.color.rgb = col; rr.font.bold = bold

doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run("Bu doküman, ekteki “MRM_Archer_İş_Analizi_Gereklilik_Kitabı.xlsx” çalışma kitabı ile birlikte kullanılır. "
                 "Anlatı ve mimari bu dokümanda; alan-seviyesi tasarım, gereklilik matrisi ve regülasyon eşlemesi Excel’dedir.")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY

doc.add_page_break()

# =========================================================
# 1. YONETICI OZETI
# =========================================================
doc.add_heading("1. Yönetici Özeti", level=1)
para("Banka, kredi, piyasa, operasyonel ve karşı taraf risk modellerinden TFRS 9 beklenen kredi zararı, İSEDES "
     "sermaye, ALM/IRRBB, AML/suistimal, fiyatlama-değerleme modellerine ve giderek artan yapay zekâ (AI/ML) ve "
     "üretken yapay zekâ (GenAI) uygulamalarına kadar geniş bir model portföyü işletmektedir. Bu modellerin hatalı "
     "veya amaç dışı kullanımından doğan “model riski”, uluslararası düzenlemelerde artık kendi başına yönetilmesi "
     "gereken bir risk türü olarak ele alınmaktadır.")
para("Bu doküman; bankanın tüm model ve algoritmalarının tek bir kurumsal envanterde toplanması, kritikliğe göre "
     "derecelendirilmesi, bağımsız validasyonu, sürekli izlenmesi, etki analizleri ve yönetişiminin uçtan uca bir "
     "süreç olarak mevcut Archer platformu üzerinde tesis edilmesi için gereklilikleri ve çözüm tasarımını tanımlar.")
para("Tasarım, dört uluslararası çerçevenin ortak omurgasına (US Fed SR 11-7, OSFI E-23:2027, PRA SS1/23, ECB TRIM) "
     "ve AI/ML için EU AI Act + NIST AI RMF gerekliliklerine dayandırılmıştır. Türkiye’de (BDDK) bugün müstakil bir "
     "model riski yönetimi düzenlemesi bulunmamakta; ancak model, validasyon ve backtesting yükümlülükleri "
     "İSEDES/İç Sistemler Yönetmeliği, IRB/AMA validasyon rehberi (Rehber #16) ve TFRS 9 karşılık rejimine dağılmış "
     "durumdadır. Çözüm, bu mevcut yükümlülükleri bugünden karşılayacak; BDDK müstakil bir MRM düzenlemesi "
     "yayımladığında ise uyumun bir yeniden mimari değil, bir konfigürasyon/kapsam çalışması olmasını sağlayacak "
     "biçimde “ileriye dönük uyumlu” kurgulanmıştır.")
para("Çözümün belkemiği, mevcut Archer uygulamalarının yeniden kullanımıdır: yeni bir “Model Riski Yönetimi” "
     "workspace’i altında Model Envanteri, Model Validasyonu, Kritiklik Değerlendirmesi, AI/ML Ek Değerlendirmesi, "
     "Model Değişiklik ve PMA/İstisna uygulamaları kurgulanırken; Bulgular, Aksiyon Planları, Metrikler/Metrik "
     "Sonuçları, Kontroller, Risk/Risk Hiyerarşisi, Organizasyon Hiyerarşisi, Tedarikçi, Uygulamalar/Donanımlar, "
     "Bilgi Varlıkları, Yasal Mevzuatlar ve Politikalar mevcut uygulamalar cross-reference ile bağlanır.")
para("Uygulama dört fazda önerilir: F1 Temel Envanter & Yönetişim (MVP), F2 Validasyon & Bulgu Döngüsü, "
     "F3 İzleme & Raporlama, F4 AI/ML İleri & Entegrasyon.")

# =========================================================
# 2. AMAC VE KAPSAM
# =========================================================
doc.add_heading("2. Amaç ve Kapsam", level=1)
doc.add_heading("2.1 Amaç", level=2)
bullet("Bankanın tüm modellerinin ve algoritmalarının tek bir kurum geneli envanterde izlenmesi.", bold_lead="Kurumsal envanter")
bullet("Modellerin materyalite ve karmaşıklığa göre derecelendirilmesi (tiering) ve buna göre orantılı yönetim.", bold_lead="Risk-tabanlı yönetim")
bullet("Geliştirmeden bağımsız, etkin sorgulama (effective challenge) yeteneğine sahip validasyon süreci.", bold_lead="Bağımsız validasyon")
bullet("Performansın sürekli izlenmesi, etki analizleri, bulgu ve remediation döngüsü.", bold_lead="Sürekli izleme")
bullet("Bugünkü BDDK yükümlülüklerini karşılarken gelecekteki MRM düzenlemesine hazır olma.", bold_lead="İleriye dönük uyum")

doc.add_heading("2.2 Kapsam (model evreni)", level=2)
para("Kapsam, risk türünden ve teknolojiden bağımsız olarak bankanın kararlarını etkileyen tüm modelleri içerir:")
bullet("Risk türleri: Kredi/IRB, Piyasa/VaR, Operasyonel/AMA, Karşı Taraf/CVA, TFRS 9 BKZ, İSEDES sermaye, ALM/IRRBB, AML/suistimal, fiyatlama ve değerleme modelleri.")
bullet("Teknoloji sınıfları: geleneksel istatistiksel modeller, AI/ML modelleri, GenAI/temel modeller, deterministik/kural-tabanlı algoritmalar (materyal ve karmaşık olduklarında).")
bullet("Köken: iç geliştirme, satıcı (vendor) ve grup-ana modeller.")
para("Kapsam dışı: Materyal ve karmaşık olmayan basit deterministik hesaplamalar ve EUC’ler için asgari yönetim "
     "kontrolleri yeterlidir; bunlar envanterde işaretlenir ancak tam yaşam döngüsü yönetimine tabi tutulmaz.", italic=True)

# =========================================================
# 3. TANIMLAR VE KISALTMALAR
# =========================================================
doc.add_heading("3. Tanımlar ve Kısaltmalar", level=1)
add_table(
    ["Terim", "Açıklama"],
    [
     ["Model", "Girdi veriyi (nicel/nitel/uzman görüşü) istatistiksel, ekonomik, finansal veya matematiksel teknikle işleyip çıktı üreten yöntem. Üç bileşen: girdi · işleme · raporlama/çıktı."],
     ["Model riski", "Hatalı veya amaç dışı model çıktılarına dayanan kararlardan doğan olumsuz sonuç (finansal kayıp, hatalı karar, itibar zararı) potansiyeli. İki kaynak: temel hatalar; yanlış/uygunsuz kullanım."],
     ["Model envanteri", "Tüm modellerin kayıt altına alındığı, kurum geneli, güncel (evergreen) sistem-of-record."],
     ["Tier / kritiklik", "Materyalite + karmaşıklığa göre modele atanan risk derecesi; validasyon/izleme yoğunluk ve frekansını yönlendirir."],
     ["Bağımsız validasyon", "Geliştirme ve kullanımdan bağımsız tarafça modelin kavramsal sağlamlık, süreç doğrulama ve çıktı analizi ile değerlendirilmesi."],
     ["Etkin sorgulama", "Yetkin, bağımsız, teşvik ve etkiye sahip taraflarca modelin kritik analizi (effective challenge)."],
     ["Backtesting / benchmarking", "Model tahmininin gerçekleşenle (geriye dönük test) ve alternatif kaynaklarla (kıyaslama) karşılaştırılması."],
     ["PMA", "Model sonrası düzeltme / overlay / override — modelin yeterince yansıtmadığı risk için çıktıya yapılan ayarlama."],
     ["MRM", "Model Riski Yönetimi (Model Risk Management)."],
     ["İSEDES", "İçsel Sermaye Yeterliliği Değerlendirme Süreci (ICAAP)."],
     ["DDE", "Data-Driven Event (Archer’da veri-güdümlü olay/otomasyon)."],
    ],
    widths=[4.5, 13.5], font=9.5
)

# =========================================================
# 4. DUZENLEYICI CERCEVE
# =========================================================
doc.add_heading("4. Düzenleyici Çerçeve ve Referanslar", level=1)
para("Aşağıdaki çerçeveler tasarımın gereklilik kaynağıdır. Ortak omurga: model tanımı, kurumsal envanter, "
     "risk-tabanlı tiering, bağımsız validasyon + etkin sorgulama, sürekli izleme + çıktı analizi/backtesting + "
     "benchmarking, board/üst yönetim gözetimi, kapsamlı dokümantasyon ve iç denetim güvencesi.")

doc.add_heading("4.1 Uluslararası çerçeveler", level=2)
bullet("Model tanımı (3 bileşen), model riskinin iki kaynağı, üç temel unsur (geliştirme; validasyon; yönetişim/politika/kontrol), "
       "‘effective challenge’ ve bağımsız validasyon, model envanteri ve risk-tabanlı yoğunluk. Prensip-bazlı, validasyon odaklı.",
       bold_lead="US Fed SR 11-7 / OCC 2011-12")
bullet("Kurum geneli (finansal olmayan modeller ve AI/ML dahil) kapsam; 3 sonuç → 12 ilke; altı-rollü yönetişim "
       "(owner, developer, reviewer, approver, user, stakeholder); prescribed asgari model envanteri şeması; "
       "altı-aşamalı yaşam döngüsü (tasarım-veri-geliştirme, review, onay, deployment, izleme, decommission); "
       "model risk rating. En net veri modeli kaynağı. Yürürlük 1 Mayıs 2027.",
       bold_lead="OSFI E-23 (2027)")
bullet("Beş ilke: (1) model tanımlama & risk sınıflandırma, (2) yönetişim (SMF hesap verebilirliği, model risk iştahı, "
       "yıllık öz-değerlendirme/attestation), (3) geliştirme-uygulama-kullanım, (4) bağımsız validasyon, "
       "(5) model riski azaltıcıları (PMA, kısıt, istisna). AI/ML ‘dinamik model’ ve karmaşıklık faktörleriyle kapsanır. "
       "Yürürlük 17 Mayıs 2024.",
       bold_lead="PRA SS1/23 (Bank of England)")
bullet("Düzenleyici (IRB/piyasa/karşı taraf) modeller için bağımsız validasyon birimi, yıllık validasyon, backtesting/"
       "discriminatory power/representativeness/override/stability/benchmarking test kataloğu ve eşikleri; 7+ boyutlu "
       "veri kalitesi çerçevesi (DQF). En net validasyon-testi ve veri-kalitesi kaynağı.",
       bold_lead="ECB TRIM / Guide to Internal Models")

doc.add_heading("4.2 AI/ML için ek çerçeveler", level=2)
bullet("Kredi değerliliği/skorlama AI’ı ‘yüksek riskli’ (Annex III 5b); Art.9-15 (risk yönetimi, veri yönetişimi, "
       "teknik dokümantasyon, loglama, şeffaflık, insan gözetimi, robustluk/güvenlik) + Art.27 FRIA. Yüksek-risk kredi "
       "yükümlülükleri 2 Ağustos 2026’dan itibaren.", bold_lead="EU AI Act (2024/1689)")
bullet("Govern / Map / Measure / Manage fonksiyonları ve yedi güvenilir-AI özelliği (geçerli-güvenilir, güvenli, "
       "güvenli-dayanıklı, hesap verebilir-şeffaf, açıklanabilir, gizlilik, adil). GenAI Profili (AI 600-1) ek overlay.",
       bold_lead="NIST AI RMF (AI 100-1)")
bullet("Açıklanabilirlik, yanlılık/fairness, veri lineage, robustluk, reproducibility, drift, retraining yönetişimi, "
       "insan gözetimi, temel model/3. taraf bağımlılığı ve GenAI guardrail/halüsinasyon izleme — geleneksel modele EK gereklilikler.",
       bold_lead="Sektör çerçeveleri (EY, WF GenAI, EBA ML-IRB, FSB, MAS FEAT, HKMA)")

doc.add_heading("4.3 Türkiye (BDDK) — mevcut yükümlülükler ve ileriye dönük uyum", level=2)
para("Türkiye’de müstakil bir MRM düzenlemesi bugün için yoktur ve resmi taslak aşamasında değildir. Ancak model, "
     "validasyon ve backtesting yükümlülükleri mevcut düzenlemeye dağılmıştır. Tasarım bunları bugünden karşılar:")
bullet("İç Sistemler & İSEDES Yönetmeliği (RG 11/7/2014, 29057): Md.3 ‘Validasyon’ tanımı; Md.26 iç denetimin model "
       "doğruluğu/backtesting/veri kalitesi testi; Md.42 risk ölçüm sistemleri; Md.43 stres testi; İSEDES içsel model "
       "validasyonunun geliştiriciden bağımsız ekip/uzman kuruluş tarafından yapılması ve raporların BDDK’ya sunulması.")
bullet("İçsel Derecelendirme/İleri Ölçüm Değerlendirme, Validasyon ve Kurumsal Yönetim Rehberi (Doküman 1041, #16): "
       "kavramsal sağlamlık + sürekli izleme/benchmarking + backtesting; bağımsızlık; model validasyon, sonuç onayı, "
       "model kontrolü, istisna tespiti, model güncelleme ve devre dışı bırakma. Mevcut MRM’e en yakın belge.")
bullet("TFRS 9 BKZ rejimi: Karşılıklar Yönetmeliği (RG 22/6/2016, 29750) Md.9/20 ve ECL Rehberi (Doküman 943) — "
       "ECL model süreç, veri, metodoloji ve validasyon beklentileri.")
bullet("Üç savunma hattı: iç kontrol (1.), risk yönetimi (2., bağımsız validasyonun doğal yeri), iç denetim (3.). "
       "Gelecekteki BDDK MRM kuralının, ikinci hatta ‘etkin sorgulama’ yetkili müstakil bir MRM/validasyon fonksiyonunu "
       "resmileştirmesi beklenir — organizasyon tasarımı buna göre yapılmalıdır.")
para("İleriye dönük tasarım hedefi: SR 11-7 + OSFI E-23’ü tasarım şablonu olarak alıp, BDDK’nın mevcut İSEDES/IRB-AMA/"
     "TFRS 9 kavram ve raporlama artefaktlarıyla uygulamak. Her Archer artefaktının hangi mevzuat maddesini karşıladığı "
     "bir eşleme tablosunda (Excel Sekme 14) tutulur; böylece yeni bir kurala karşı gap analizi mekanik hale gelir.", italic=True)

# =========================================================
# 5. TEMEL MRM GEREKLILIKLERI
# =========================================================
doc.add_heading("5. Temel MRM Gereklilikleri (Özet)", level=1)
para("Aşağıda temel gereklilikler ve alt detayları özetlenmiştir. Tam liste, Archer karşılığı, öncelik, kaynak "
     "düzenleme ve faz bilgisi ile birlikte Excel Sekme 1 (Gereklilik Matrisi) içindedir.")

reqs = [
 ("Model tanımı ve kapsam", [
    "Kurum genelinde tek model tanımı ve model/model-riski taksonomisi.",
    "Deterministik yöntem/EUC ayrımı: materyal+karmaşık ise MRM’in ilgili unsurları, aksi halde sağlam kontroller.",
    "AI/ML/GenAI/dinamik model bayrakları ve tüm risk türlerinin kapsanması."]),
 ("Model envanteri", [
    "Kurum geneli tekil, doğru, güncel envanter (kullanımda + geliştirmede + emekli).",
    "OSFI E-23 asgari alan seti + SS1/23 envanter nitelikleri (amaç, sınırlamalar, validasyon bulguları, yönetişim).",
    "Model bağımlılıkları (upstream/downstream) ve satıcı/grup modeli ayrımı."]),
 ("Kritiklik / tiering", [
    "Tutarlı, kurum geneli materyalite + karmaşıklık derecelendirmesi (skorlu anket).",
    "İçsel risk = zafiyet × materyalite; artık risk ayrı raporlanır.",
    "Tier, validasyon ve izleme derinlik/frekansını yönlendirir (frekans banka politikasıdır); periyodik bağımsız gözden geçirme."]),
 ("Geliştirme, uygulama, kullanım", [
    "Amaç/tasarım beyanı, kavramsal sağlamlık, veri uygunluğu ve temsil.",
    "Geliştirme testleri (geriye/ileriye dönük + duyarlılık), model düzeltmeleri/uzman görüşünün kaydı.",
    "Bağımsız uzmanın yeniden üretebileceği dokümantasyon; test edilmiş ve değişiklik-kontrollü uygulama ortamı."]),
 ("Bağımsız validasyon", [
    "Geliştirmeden bağımsız fonksiyon; kavramsal sağlamlık + süreç doğrulama + performans izleme/çıktı analizi.",
    "Periyodik revalidasyon; validasyon ÖNERİSİ ile ONAY kararının ayrı olması.",
    "Etkin sorgulama; bulguların dokümante edilip izlenmesi."]),
 ("Sürekli izleme", [
    "Metrik seti ve eşikler (AUC/Gini/KS, kalibrasyon, PSI/CSI drift, fairness, override).",
    "Backtesting ve benchmarking; drift/bozulma eşiği → revalidasyon/retrain tetiği."]),
 ("Model riski azaltıcıları", [
    "PMA/overlay kurum geneli tutarlı süreç, bağımsız gözden geçirme, PMA’lı/PMA’sız raporlama.",
    "Kullanım kısıtları/limitleri; istisna tanımı ve resmi eskalasyon."]),
 ("Yönetişim", [
    "Board/üst yönetim gözetimi, model risk iştahı, sorumlu üst yönetici (SMF).",
    "MRM politikası; owner/developer/user/validator/approver rolleri; üç savunma hattı ve iç denetim.",
    "Yıllık öz-değerlendirme/attestation ve finansal raporlama için MRM etkinlik raporu."]),
 ("AI/ML ek gereklilikleri", [
    "Açıklanabilirlik, yanlılık/fairness, veri lineage, robustluk/reproducibility/güvenlik.",
    "Retraining yönetişimi, insan gözetimi, etik/FRIA, temel model/3. taraf bağımlılığı.",
    "GenAI guardrail ve halüsinasyon/toksisite izleme; EU AI Act risk sınıflandırması."]),
]
for h, items in reqs:
    doc.add_heading("5." + str(reqs.index((h, items))+1) + " " + h, level=2)
    for it in items:
        bullet(it)

# =========================================================
# 6. ARCHER COZUM MIMARISI
# =========================================================
doc.add_heading("6. Önerilen Archer Çözüm Mimarisi", level=1)
para("Çözüm, mevcut Archer envanterinin (BT Risk, Bankacılık Risk, Uyumluluk, Bilgi Güvenliği, Organizasyon "
     "Yönetimi workspace’leri ve paylaşımlı Bulgu/Metrik/Kontrol uygulamaları) üzerine yeni bir “Model Riski "
     "Yönetimi (MRM)” workspace’i olarak kurgulanır. İlke: yeni uygulama yalnızca gerçekten gerektiğinde açılır; "
     "aksi halde mevcut paylaşımlı uygulamalar yeniden kullanılır.")

doc.add_heading("6.1 Yeni uygulamalar", level=2)
add_table(
    ["Uygulama", "Rol"],
    [
     ["Model Envanteri", "Bankanın tüm model/algoritmalarının ana kaydı (master); tüm ilişkilerin merkezi."],
     ["Model Kritiklik / Materyalite Değerlendirmesi", "Skorlu anket → Tier ve içsel risk skorunu üretir."],
     ["Model Validasyonu", "Bağımsız validasyon/gözden geçirme çalışmalarının kaydı; onay akışı."],
     ["AI/ML Ek Değerlendirmesi", "AI/ML/GenAI modellerine özel ek nitelikler ve kontroller (anket)."],
     ["Model Değişiklik Yönetimi", "Değişiklik/versiyon kayıtları; materyal değişiklikte revalidasyon tetiği."],
     ["PMA / İstisna Kayıtları", "Model sonrası düzeltme, kullanım kısıtı ve istisna kayıtları."],
    ],
    widths=[6.5, 11.5], font=9.5)

doc.add_heading("6.2 Yeniden kullanılan mevcut uygulamalar", level=2)
add_table(
    ["Mevcut Uygulama", "MRM’de kullanımı"],
    [
     ["Bulgular + Aksiyon Planları", "Validasyon/izleme bulguları ve remediation (‘Bulgu Kaynağı = Model Validasyonu’ değeri eklenerek)."],
     ["Metrikler / Metrik Sonuçları", "Sürekli izleme KPI/KRI’ları; eşik aşımı → Bulgu."],
     ["Risk / Risk Hiyerarşisi", "Toplulaştırılmış model riskinin kurumsal risk taksonomisine bağlanması."],
     ["Kontroller", "Model riskini azaltan kontroller."],
     ["Organizasyon – İş Hiyerarşisi", "Model sahibi birim (GMY/Daire/Bölüm)."],
     ["Tedarikçi", "Satıcı (vendor) modeller."],
     ["Uygulamalar / Donanımlar / IT Hizmetleri", "Modelin koştuğu sistem ve altyapı."],
     ["Bilgi Varlıkları Envanteri", "Model girdi verisi / veri varlığı ve lineage."],
     ["Yasal Mevzuatlar / Politikalar ve Standartlar", "Regülasyon eşleme ve MRM politikası/DQF."],
     ["Risk Kabul Talepleri / Denetlenen Görüşleri", "İstisna/kabul akışı ve iç denetim (3. hat) gözetimi."],
    ],
    widths=[6.5, 11.5], font=9.5)
para("İlişki haritasının tamamı Excel Sekme 10’dadır.", italic=True)

doc.add_heading("6.3 Archer teknik mekanizmaları", level=2)
bullet("Seçim listeleri (Values Lists): model tipi, kategori, yaşam döngüsü durumu, tier, validasyon türü/sonucu (Excel Sekme 9).")
bullet("Gelişmiş iş akışı (Advanced Workflow): Geliştirici → Bağımsız Validasyon → Model Risk Komitesi → Onay/Ret/Koşullu.")
bullet("Hesaplanan alanlar: sonraki validasyon tarihi (= son + tier frekansı), içsel risk skoru, gecikme bayrağı.")
bullet("Veri-güdümlü olaylar (DDE): durum geçişleri, koşullu alan görünürlüğü/zorunluluğu, validasyon başarısızsa otomatik bulgu.")
bullet("Bildirimler: yaklaşan/geciken revalidasyon, izleme eşik aşımı, açık bulgu SLA, PMA kaldırma tetiği.")
bullet("Erişim rolleri: 1./2./3. hat ayrımı; kayıt/alan seviyesi güvenlik ile validasyon bağımsızlığının sistemsel garantisi.")
bullet("Veri beslemeleri (Data Feeds): MLOps/model geliştirme ortamından envanter ve metrik beslemesi (F4, opsiyonel).")

# =========================================================
# 7. YASAM DONGUSU VE IS AKISLARI
# =========================================================
doc.add_heading("7. Model Yaşam Döngüsü ve İş Akışları", level=1)
para("Model kaydı bir durum makinesi olarak yönetilir. Aşağıdaki akış Advanced Workflow ve DDE ile uygulanır; "
     "tam geçiş kuralları Excel Sekme 12’dedir.")
para("Taslak/Tanımlama → Geliştirme → Bağımsız Validasyonda → Onayda → Onaylı/Prod’da → İzlemede → (Revalidasyonda) "
     "→ … ; yan durumlar: Koşullu Onaylı (PMA/kısıt ile), Kısıtlı/Askıda, Emekli/Devre Dışı.", size=11)
add_table(
    ["Durum", "Sonraki", "Geçiş koşulu"],
    [
     ["Geliştirme", "Bağımsız Validasyonda", "Geliştirme testleri + dokümantasyon tamam"],
     ["Bağımsız Validasyonda", "Onayda / Geliştirme", "Validasyon raporu + öneri hazır"],
     ["Onayda", "Onaylı/Prod · Koşullu · Reddedildi", "Onay makamı kararı (validasyondan ayrı)"],
     ["Onaylı/Prod", "İzlemede", "Deployment testleri tamam"],
     ["İzlemede", "Revalidasyonda · Kısıtlı · Değişiklik", "Frekans dolumu / eşik aşımı / değişiklik"],
     ["Kısıtlı/Askıda", "İzlemede · Emekli", "Ciddi eksik giderildi / karar"],
     ["Emekli", "(son)", "Ciddi hata / tekrar ihlal / obsolescence"],
    ],
    widths=[5.5, 6.5, 6.0], font=9.5)

# =========================================================
# 8. UYGULAMA ALAN TASARIMI OZETI
# =========================================================
doc.add_heading("8. Uygulama Bazında Alan Tasarımı (Özet)", level=1)
para("Her uygulamanın tam alan tasarımı (alan adı, Archer alan tipi, seçim listesi, zorunluluk, hesaplama/çapraz "
     "referans, açıklama) Excel’de ilgili sekmededir. Öne çıkanlar:")
bullet("Model Envanteri (Sekme 2): kimlik, sınıflandırma/bayraklar, köken, amaç/tasarım, kritiklik (Tier + skorlar), roller, "
       "bağımlılıklar, sistem/veri linkleri, yaşam döngüsü, validasyon takibi, kısıt, dokümantasyon, saklama.")
bullet("Model Validasyonu (Sekme 3): tür, kapsam, bağımsızlık teyidi, kavramsal sağlamlık, süreç doğrulama, çıktı analizi, "
       "etkin sorgulama, validasyon ÖNERİSİ ve ONAY kararı (ayrı alanlar), rapor.")
bullet("Kritiklik/Tiering (Sekme 4): materyalite + karmaşıklık + bağımlılık/otonomi soruları, ağırlıklar ve tier eşleme tablosu.")
bullet("AI/ML Ek (Sekme 5): algoritma ailesi, temel model/3. taraf, veri/eğitim, açıklanabilirlik, fairness, robustluk, "
       "izleme/drift, retraining, insan gözetimi, etik/FRIA, EU AI Act risk kategorisi.")
bullet("PMA/İstisna (Sekme 7) ve Model Değişiklik (Sekme 8): azaltıcılar ve değişiklik/revalidasyon tetikleri.")

# =========================================================
# 9. ROLLER VE SORUMLULUKLAR
# =========================================================
doc.add_heading("9. Roller ve Sorumluluklar (Üç Savunma Hattı)", level=1)
bullet("Model sahibi, geliştirici ve kullanıcı; modeli geliştirir, kullanır ve günlük performansı izler.", bold_lead="1. hat")
bullet("MRM fonksiyonu ve bağımsız validasyon; çerçeve, kontroller, etkin sorgulama ve onay önerisi.", bold_lead="2. hat")
bullet("İç denetim; MRM çerçevesinin etkinliğini ve validasyon fonksiyonunu bağımsızca denetler.", bold_lead="3. hat")
para("Onay yetkisi, validasyondan ayrı olarak Model Risk Komitesi/üst yönetimdedir. Ayrıntılı RACI matrisi Excel "
     "Sekme 11’dedir.")

# =========================================================
# 10. RAPORLAMA
# =========================================================
doc.add_heading("10. Raporlama ve Dashboard’lar", level=1)
bullet("Model envanteri kırılımı (tier/tip/kategori/durum).")
bullet("Validasyon takvimi ve gecikmeler; açık model bulguları ve remediation durumu.")
bullet("Toplulaştırılmış model riski ısı haritası; iştaha karşı profil (board paneli).")
bullet("AI/ML model portföyü ve drift/izleme durumu; Model Risk Komitesi paneli.")
bullet("BDDK İçsel Model Validasyon raporu ve İSEDES model bölümlerinin talep üzerine üretimi.")

# =========================================================
# 11. YOL HARITASI
# =========================================================
doc.add_heading("11. Uygulama Yol Haritası", level=1)
add_table(
    ["Faz", "Kapsam", "Temel çıktı"],
    [
     ["F1 – Temel Envanter & Yönetişim (MVP)", "Model Envanteri, tiering anketi, roller, durum makinesi, seçim listeleri, erişim rolleri", "Kurum geneli envanter + kritiklik + MRM politikası + board paneli iskeleti"],
     ["F2 – Validasyon & Bulgu Döngüsü", "Model Validasyonu, Bulgular entegrasyonu, geliştirme alanları, bağımlılıklar, satıcı model", "Bağımsız validasyon iş akışı + validasyon→bulgu→aksiyon + BDDK raporu"],
     ["F3 – İzleme & Raporlama", "Metrik seti, dashboard’lar, toplulaştırılmış raporlama, öz-değerlendirme", "İzleme metrik kataloğu + risk ısı haritası + yıllık öz-değerlendirme"],
     ["F4 – AI/ML İleri & Entegrasyon", "AI/ML derinleştirme, GenAI guardrail, DQF/BCBS239, data feeds", "GenAI izleme + DQF + MLOps beslemesi (opsiyonel)"],
    ],
    widths=[4.8, 7.7, 5.5], font=9)

# =========================================================
# 12. VARSAYIMLAR, BAGIMLILIKLAR, RISKLER
# =========================================================
doc.add_heading("12. Varsayımlar, Bağımlılıklar ve Riskler", level=1)
doc.add_heading("12.1 Varsayımlar", level=2)
bullet("Mevcut Archer uygulamaları (Bulgular, Metrikler, Risk, Kontroller, Organizasyon, Tedarikçi, Uygulamalar, Bilgi Varlıkları) erişilebilir ve cross-reference’a uygundur.")
bullet("Tier→frekans eşlemesi banka MRM politikasınca belirlenecek bir parametredir (regülasyon sabit sayı dayatmaz).")
bullet("Bağımsız validasyon fonksiyonu ikinci hatta organizasyonel olarak konumlandırılabilir.")
doc.add_heading("12.2 Bağımlılıklar", level=2)
bullet("MRM politikası, model risk iştahı ve tiering metodolojisinin iş tarafınca onaylanması (F1 önkoşulu).")
bullet("MLOps/model geliştirme ortamıyla veri beslemesi entegrasyonu (F4, opsiyonel).")
doc.add_heading("12.3 Riskler", level=2)
bullet("Kapsam genişliği: tüm risk türleri + AI/ML aynı anda; fazlama ile yönetilir (MVP → tam süreç).")
bullet("Veri kalitesi ve envanter güncelliği (evergreen) olmadan raporlama güvenilirliği düşer.")
bullet("BDDK düzenlemesi yayımlandığında bazı alanların yeniden eşlenmesi gerekebilir; regülasyon eşleme tablosu bunu mekanik kılar.")

# =========================================================
# 13. EKLER
# =========================================================
doc.add_heading("13. Ekler", level=1)
bullet("Ek-A: MRM_Archer_İş_Analizi_Gereklilik_Kitabı.xlsx — 16 sekme (gereklilik matrisi, alan tasarımları, metrik kataloğu, values list’ler, ilişki haritası, RACI, yaşam döngüsü, yol haritası, regülasyon eşleme, referanslar).")
bullet("Ek-B: Birincil kaynak URL’leri — Excel Sekme 15.")

para("")
para("Not: Bu taslak, regülasyon araştırmasına dayalı bir başlangıç setidir; bankanın mevcut model politikası, "
     "organizasyon yapısı ve Archer lisans/modül kapsamına göre birlikte netleştirilerek nihai hale getirilecektir.",
     italic=True, color=GREY, size=9.5)

out = "/home/user/rrr/mrm-archer/MRM_Archer_Is_Analizi_ve_Cozum_Tasarimi.docx"
doc.save(out)
print("KAYDEDILDI:", out)
